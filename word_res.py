import pandas as pd
from docx import Document
from docx.shared import Inches
import os

def generate_word_report_from_csv_folder(folder_path, output_docx_path="model_results_summary.docx"):
    # Create a new Word document
    doc = Document()
    doc.add_heading('Model Results Summary', level=1)

    # Define columns to exclude (lowercase, stripped)
    exclude_columns = ['last_fold_preds', 'last_fold_true']

    # Get all CSV files in the folder
    csv_files = [f for f in os.listdir(folder_path) if f.endswith('.csv')]
    if not csv_files:
        print("No CSV files found in the specified folder.")
        return

    for file_name in sorted(csv_files):
        file_path = os.path.join(folder_path, file_name)
        try:
            df = pd.read_csv(file_path)
        except Exception as e:
            doc.add_heading(f"{file_name} (Error Reading File)", level=2)
            doc.add_paragraph(str(e))
            continue

        doc.add_heading(file_name.replace(".csv", ""), level=2)

        # Strip whitespace from column names
        df.columns = [col.strip() for col in df.columns]

        # Drop columns that match the exclusion list (case-insensitive)
        df = df[[col for col in df.columns if col.strip().lower() not in exclude_columns]]

        if df.empty:
            doc.add_paragraph("No data available after removing excluded columns.")
            continue

        # Round all numeric values to 3 decimal places
        df = df.applymap(lambda x: round(x, 3) if isinstance(x, (int, float)) else x)

        # Create table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

    # Save the document
    doc.save(output_docx_path)
    print(f"Word document created: {output_docx_path}")

# Example usage
generate_word_report_from_csv_folder("last_results", "output_summary.docx")
